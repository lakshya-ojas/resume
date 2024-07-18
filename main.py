import PyPDF2
import docx
import spacy
from spacy.lang.en.stop_words import STOP_WORDS
from spacy.matcher import Matcher, PhraseMatcher
from spacy.lang.en import English
from gradio_client import Client
import re
import json
import os
import sys
from database_handler import create_table, insert_data
from summarize_chatgpt import get_summery, get_summery2
import requests

# Load spaCy model
nlp = spacy.load('en_core_web_lg')

def convert_pdf_to_text(pdf_path):
    """Converts a PDF file to text."""
    with open(pdf_path, 'rb') as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        num_pages = len(pdf_reader.pages)
        text = ''
        for page_num in range(num_pages):
            page = pdf_reader.pages[page_num]
            text += str (page.extract_text())
    return text

def convert_docx_to_text(docx_path):
    """Converts a DOCX file to text."""
    doc = docx.Document(docx_path)
    docx_text = '\n'.join([str(para.text) for para in doc.paragraphs])
    return docx_text

#work on it
def extract_text_from_doc(doc_file):
    if os.path.exists(doc_file):
        doc = docx.Document(doc_file)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(str(para.text))
        return '\n'.join(full_text)
    else:
        raise FileNotFoundError(f"File '{doc_file}' not found.")


def convert_txt_to_text(txt_path):
    """Converts a TXT file to text."""
    with open(txt_path, 'r', encoding='utf-8') as txt_file:
        txt_text = txt_file.read()
    return txt_text 

def save_text_to_file(text, output_path):
    """Saves text to a file."""
    with open(output_path, 'w', encoding='utf-8') as output_file:
        output_file.write(text)

def text_summarizer(text):
    """Summarizes text."""
    doc = nlp(text)
    stopwords = list(STOP_WORDS)
    word_freq = {}
    for word in doc:
        if word.text.lower() not in stopwords:
            if word.text.lower() not in word_freq:
                word_freq[word.text.lower()] = 1
            else:
                word_freq[word.text.lower()] += 1

    max_freq = max(word_freq.values())
    for word in word_freq:
        word_freq[word] = (word_freq[word] / max_freq)

    # for ent in doc.ents:
    #     if ent.label_ == "PERSON":
    #         print(f"Person: {ent.text}")
    #     elif ent.label_ == "DATE":
    #         print(f"Date: {ent.text}")
    #     elif ent.label_ == "GPE" or ent.label_ == "LOC":
    #         print(f"Location: {ent.text}")

    sentence_list = [sentence for sentence in doc.sents]
    sentence_scores = {}
    for sent in sentence_list:
        for word in sent:
            if word.text.lower() in word_freq:
                if sent not in sentence_scores:
                    sentence_scores[sent] = word_freq[word.text.lower()]
                else:
                    sentence_scores[sent] += word_freq[word.text.lower()]

    from heapq import nlargest
    summarized_sentences = nlargest(3, sentence_scores, key=sentence_scores.get)
    summarized_text = ' '.join([w.text for w in summarized_sentences])
    return summarized_text


def text_summarizer2(text):
    """Summarizes text using BERT Summary."""
    API_URL = "https://api-inference.huggingface.co/models/Shobhank-iiitdwd/BERT_summary"
    headers = {"Authorization": "Bearer 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"}

    def query(payload):
        response = requests.post(API_URL, headers=headers, json=payload)
        return response.json()
        
    output = query({
        "inputs": text,
    })

    print("this is the text_summarizer2")
    print(output)
    
    # Check if output is a list and extract summary_text if present
    if isinstance(output, list) and len(output) > 0 and 'summary_text' in output[0]:
        return output[0]['summary_text']
    else:
        return "Summary not available"  # Handle the case where summary_text is not found


def text_summarizer3(text):
    """Summarizes text using ChatGPT."""
    
    client = Client("huggingface-projects/llama-2-13b-chat")
    new_text = "summary of the text: " + text + "\n" + " in utf-8 fromate only"
    result = client.predict(
            message=new_text,
            system_prompt="Hello!!",
            max_new_tokens=1024,
            temperature=0.6,
            top_p=0.9,
            top_k=50,
            repetition_penalty=1.2,
            api_name="/chat"
    )
    print(result)


def extract_entities(text):
    """Extracts entities (persons, phone numbers, emails) from text."""
    phone_pattern = r'\b(?:0|\+?91)?[789]\d{9}\b'  # Matches Indian mobile numbers
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    doc = nlp(text)
    
    entities = {
        "persons": [ent.text for ent in doc.ents if ent.label_ == "PERSON"],
        "phone_numbers": re.findall(phone_pattern, text),
        "emails": re.findall(email_pattern, text)
    }
    return entities

def convert_file_to_text(input_path, output_path):
    """Converts various file formats to text."""
    text = ""
    if input_path.endswith('.pdf'):
        text = convert_pdf_to_text(input_path)
    elif input_path.endswith('.docx'):
        text = convert_docx_to_text(input_path)
    elif input_path.endswith('.txt'):
        text = convert_txt_to_text(input_path)
    # elif input_path.endswith(".doc"):
    #     text = extract_text_from_doc(input_path)
    else:
        raise ValueError("Unsupported file type. Only PDF, DOCX, and TXT files are supported.")
    
    print("Success in extracting the text from the file")
    save_text_to_file(text, "./output/ExtractedText.txt")

    summarized_text = text_summarizer(text)
    # summarized_text = text_summarizer2(text)
    # summarized_text = text_summarizer3(text)
    # summarized_text = get_summery(text)
    # summarized_text = get_summery2(text)

    print("Summarized text is done using spaCy")

    save_text_to_file(summarized_text, "./output/summarize_text.txt")
    print("File is saved in the output folder for summarize and extracted text")

    entities = extract_entities(text)
    print("User information entities are extracted from the text")

    return text, summarized_text, entities
